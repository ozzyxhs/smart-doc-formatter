"""检查 docx 真实排版落地：边距 / eastAsia 分绑 / 标题字体 / 三线表 / 页眉页脚。
用法: python scripts/inspect_docx.py <docx>
"""
import sys
import io
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
from docx import Document
from docx.oxml.ns import qn

path = sys.argv[1] if len(sys.argv) > 1 else r"app\_jobs\smoke_output.docx"
d = Document(path)


def mm(v):
    return round(v / 36000, 1) if v else None


s = d.sections[0]
print("边距 T/B/L/R mm:", mm(s.top_margin), mm(s.bottom_margin), mm(s.left_margin), mm(s.right_margin))
print("页面 mm:", mm(s.page_width), "x", mm(s.page_height))
sectPr = s._sectPr
dg = sectPr.find(qn("w:docGrid"))
print("docGrid:", dg.get(qn("w:type")) if dg is not None else None)

# 页眉
hp = s.header.paragraphs[0]
pBdr = hp._p.find(qn("w:pPr") + "/" + qn("w:pBdr")) if hp._p.find(qn("w:pPr")) is not None else None
bottom = None
pPr = hp._p.find(qn("w:pPr"))
if pPr is not None:
    b = pPr.find(qn("w:pBdr"))
    if b is not None:
        bb = b.find(qn("w:bottom"))
        if bb is not None:
            bottom = (bb.get(qn("w:val")), bb.get(qn("w:sz")))
print("页眉文字:", hp.text[:40], "| 双线 bottom(val,sz):", bottom)

# 页脚 PAGE 域
fp = s.footer.paragraphs[0]
has_page = any(e.text == "PAGE" for e in fp._p.iter(qn("w:instrText")))
print("页脚 文本:", repr(fp.text), "| 含PAGE域:", has_page)


def font_of(run):
    rPr = run._element.find(qn("w:rPr"))
    ea = lat = sz = None
    if rPr is not None:
        rf = rPr.find(qn("w:rFonts"))
        if rf is not None:
            ea = rf.get(qn("w:eastAsia")); lat = rf.get(qn("w:ascii"))
        szel = rPr.find(qn("w:sz"))
        if szel is not None:
            sz = int(szel.get(qn("w:val"))) / 2
    return ea, lat, sz, run.bold


print("\n--- 段落抽样 (eastAsia, latin, pt, bold) ---")
shown = 0
for p in d.paragraphs:
    if not p.text.strip() or not p.runs:
        continue
    ea, lat, sz, b = font_of(p.runs[0])
    print(f"   ea={ea} lat={lat} sz={sz} b={b} | {p.text[:34]}")
    shown += 1
    if shown >= 14:
        break

# 三线表
if d.tables:
    t = d.tables[0]
    tblPr = t._tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    res = {}
    if borders is not None:
        for edge in ("top", "bottom", "insideV", "insideH"):
            e = borders.find(qn(f"w:{edge}"))
            if e is not None:
                res[edge] = (e.get(qn("w:val")), e.get(qn("w:sz")))
    print("\n三线表 borders:", res, "| 表数:", len(d.tables))
