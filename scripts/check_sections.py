# -*- coding: utf-8 -*-
"""检查分节页码方案：每节的 pgNumType(fmt/start) + 页眉文字 + 页脚是否有 PAGE 域。"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
from docx import Document
from docx.oxml.ns import qn

path = sys.argv[1] if len(sys.argv) > 1 else r"app\_jobs\smoke_output.docx"
d = Document(path)
print("节数", len(d.sections))
for i, s in enumerate(d.sections):
    pg = s._sectPr.find(qn("w:pgNumType"))
    fmt = pg.get(qn("w:fmt")) if pg is not None else None
    start = pg.get(qn("w:start")) if pg is not None else None
    hp = s.header.paragraphs[0].text.strip()[:18]
    fp_has_page = any(e.text == "PAGE" for e in s.footer.paragraphs[0]._p.iter(qn("w:instrText")))
    print(f"节{i}: pgNumFmt={fmt} start={start} | 页眉=\"{hp}\" | 页脚PAGE域={fp_has_page}")
