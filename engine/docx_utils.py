"""python-docx + oxml 补丁工具：中文 Word 排版的硬骨头集中在这里。

eastAsia/latin 分绑、字符缩进、行单位段间距、页眉粗细双线、三线表、文档网格、页码域。
"""
import re
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

ALIGN = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def pt_of(template, size_name, default=10.5):
    """字号 -> pt。健壮：数字 / 中文字号名(小二…) / 'Npt'·'N磅'·'Npx'·'N' 字符串。
    认不出兜底 default（默认五号 10.5），绝不抛错——编译/聊天产出的模板字号格式可能五花八门。
    """
    if isinstance(size_name, (int, float)):
        return float(size_name)
    if size_name is None:
        return default
    s = str(size_name).strip()
    tbl = template.get("size_table", {})
    if s in tbl:
        return float(tbl[s])
    m = re.match(r"^([0-9]+(?:\.[0-9]+)?)\s*(?:pt|px|磅)?$", s)
    if m:
        return float(m.group(1))
    return default


def _fontname(v, default=None):
    """字体名强转：str 原样；dict 取 cn/name/第一个字符串值；其余 default。防 dict 喂进 oxml 崩。"""
    if isinstance(v, str):
        return v
    if isinstance(v, dict):
        for k in ("cn", "name", "latin", "default_cn"):
            if isinstance(v.get(k), str):
                return v[k]
        for vv in v.values():
            if isinstance(vv, str):
                return vv
    return default


def _num(v, default=0.0):
    """数值强转：数字原样；'25.4mm'/'3pt' 取数；dict/认不出 -> default。"""
    if isinstance(v, (int, float)):
        return float(v)
    m = re.match(r"^\s*([0-9]+(?:\.[0-9]+)?)", str(v)) if v is not None else None
    return float(m.group(1)) if m else default


def _get_or_add(parent, tag):
    el = parent.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        parent.append(el)
    return el


def set_run_font(run, *, cn=None, latin=None, size_pt=None, bold=None, italic=None):
    """关键：eastAsia(中文) 与 latin(西文) 分开绑定。font 名/字号强转，防 dict 喂进 oxml 崩。"""
    cn = _fontname(cn)
    latin = _fontname(latin)
    if size_pt is not None:
        sp = _num(size_pt, 0)
        if sp:
            run.font.size = Pt(sp)
    if bold is not None:
        run.font.bold = bool(bold)
    if italic is not None:
        run.font.italic = bool(italic)
    rPr = run._element.get_or_add_rPr()
    rFonts = _get_or_add(rPr, "w:rFonts")
    if latin:
        rFonts.set(qn("w:ascii"), latin)
        rFonts.set(qn("w:hAnsi"), latin)
    if cn:
        rFonts.set(qn("w:eastAsia"), cn)


def set_paragraph_format(p, *, align=None, first_line_chars=None, line_spacing_single=False,
                         before_lines=None, after_lines=None, page_break_before=None):
    pf = p.paragraph_format
    if align in ALIGN:
        p.alignment = ALIGN[align]
    if line_spacing_single:
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if page_break_before is not None:
        pf.page_break_before = page_break_before
    pPr = p._p.get_or_add_pPr()
    if first_line_chars is not None:
        ind = _get_or_add(pPr, "w:ind")
        ind.set(qn("w:firstLineChars"), str(int(_num(first_line_chars, 0) * 100)))
        ind.set(qn("w:firstLine"), "0")     # 让 chars 生效
    if before_lines is not None or after_lines is not None:
        sp = _get_or_add(pPr, "w:spacing")
        if before_lines is not None:
            sp.set(qn("w:beforeLines"), str(int(_num(before_lines, 0) * 100)))
        if after_lines is not None:
            sp.set(qn("w:afterLines"), str(int(_num(after_lines, 0) * 100)))


def set_page(section, template):
    pg = template.get("page", {})
    section.page_width = Mm(_num(pg.get("width_mm"), 210))
    section.page_height = Mm(_num(pg.get("height_mm"), 297))
    m = pg.get("margins_mm", {}) if isinstance(pg.get("margins_mm"), dict) else {}
    section.top_margin = Mm(_num(m.get("top"), 25.4))
    section.bottom_margin = Mm(_num(m.get("bottom"), 25.4))
    section.left_margin = Mm(_num(m.get("left"), 25.4))
    section.right_margin = Mm(_num(m.get("right"), 25.4))
    section.header_distance = Mm(_num(template.get("header", {}).get("margin_mm"), 15))
    section.footer_distance = Mm(_num(template.get("footer", {}).get("margin_mm"), 15))
    # 文档网格 38 行 × 38 字
    grid = pg.get("grid")
    if grid:
        sectPr = section._sectPr
        dg = _get_or_add(sectPr, "w:docGrid")
        dg.set(qn("w:type"), "linesAndChars")
        dg.set(qn("w:linePitch"), "312")            # ~ 行距
        dg.set(qn("w:charSpace"), "0")


def set_double_bottom_border(paragraph, upper_pt=3):
    """页眉粗细双线（粗线在上）。sz 单位 = 1/8 pt。"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = _get_or_add(pPr, "w:pBdr")
    bottom = _get_or_add(pBdr, "w:bottom")
    bottom.set(qn("w:val"), "thickThinSmallGap")    # 粗在上、细在下
    bottom.set(qn("w:sz"), str(int(_num(upper_pt, 3) * 8)))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")


def add_page_number_field(paragraph, fmt="- n -"):
    """在段落里插入 { PAGE } 域，并按 fmt 包裹前后缀（如 '- n -'）。"""
    pre, post = "", ""
    if "n" in fmt:
        pre, post = fmt.split("n", 1)

    def _run_text(t):
        r = OxmlElement("w:r")
        el = OxmlElement("w:t")
        el.set(qn("xml:space"), "preserve")
        el.text = t
        r.append(el)
        return r

    if pre:
        paragraph._p.append(_run_text(pre))
    # field begin
    r1 = OxmlElement("w:r"); f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin"); r1.append(f1)
    r2 = OxmlElement("w:r"); instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"; r2.append(instr)
    r3 = OxmlElement("w:r"); f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end"); r3.append(f3)
    paragraph._p.append(r1); paragraph._p.append(r2); paragraph._p.append(r3)
    if post:
        paragraph._p.append(_run_text(post))


def _clear_runs(paragraph):
    for r in list(paragraph.runs):
        r._element.getparent().remove(r._element)


def set_page_number_format(section, numfmt, start=1):
    """分节页码：numfmt = decimal | lowerRoman | upperRoman；start 重起页号。"""
    sectPr = section._sectPr
    pg = _get_or_add(sectPr, "w:pgNumType")
    pg.set(qn("w:fmt"), numfmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))


def setup_title_header(section, template, title):
    """本节页眉 = 论文题目 + 粗细双线（封面节不要调用）。"""
    h = template["header"]
    section.header.is_linked_to_previous = False
    p = section.header.paragraphs[0]
    _clear_runs(p)
    run = p.add_run(title or "")
    set_run_font(run, cn=h["font"]["cn"], latin=h["font"]["latin"],
                 size_pt=pt_of(template, h["font"]["size"]))
    set_paragraph_format(p, align=h["font"]["align"])
    set_double_bottom_border(p, upper_pt=h["border_below"]["upper_pt"])


def setup_pagenum_footer(section, template, fmt_str):
    """本节页脚 = 页码域（数字格式随本节 pgNumType）。"""
    f = template["footer"]
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    _clear_runs(p)
    set_paragraph_format(p, align=f["font"]["align"])
    add_page_number_field(p, fmt=fmt_str)
    for r in p.runs:
        set_run_font(r, cn=f["font"]["cn"], latin=f["font"]["latin"],
                     size_pt=pt_of(template, f["font"]["size"]))


def blank_header_footer(section):
    """封面/扉页节：无页眉、无页码。"""
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    _clear_runs(section.header.paragraphs[0])
    _clear_runs(section.footer.paragraphs[0])


def apply_three_line_table(table, *, top_bottom_pt=1.5, middle_pt=0.5):
    """开放式三线格：仅上沿/下沿(1.5pt) + 表头下(0.5pt)，无竖线无内线。"""
    tblPr = table._tbl.tblPr
    # 去掉表格自带样式
    borders = _get_or_add(tblPr, "w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = _get_or_add(borders, f"w:{edge}")
        if edge == "top" or edge == "bottom":
            e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(int(top_bottom_pt * 8)))
        else:
            e.set(qn("w:val"), "none"); e.set(qn("w:sz"), "0")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), "000000")
    # 表头行底线 0.5pt
    if table.rows:
        hdr = table.rows[0]
        for cell in hdr.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcb = _get_or_add(tcPr, "w:tcBorders")
            b = _get_or_add(tcb, "w:bottom")
            b.set(qn("w:val"), "single"); b.set(qn("w:sz"), str(int(middle_pt * 8)))
            b.set(qn("w:space"), "0"); b.set(qn("w:color"), "000000")
