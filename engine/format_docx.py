"""确定性排版：读 区块流 + 结构标签 + 模板 YAML，重建一份合规 docx。

核心保证：**文本逐 run 原样拷贝**（内容守恒 by construction），只接管排版层。
分三节：封面(无页眉无码) / 前置(摘要·目录, 罗马页码) / 正文(阿拉伯, 重起1)。
"""
import re
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Mm, Pt
from . import docx_utils as DU
from . import config

_FRONT_MARKERS = {"abstract_cn_title", "abstract_en_title", "abstract_cn_body",
                  "abstract_en_body", "keywords_cn", "keywords_en", "toc_title", "toc_item"}
_FIELD_KW = ("学院", "专业", "班级", "姓名", "学号", "指导教师", "研究方向", "导师")


def _opt_pt(template, size):
    """字号有值才转 pt，缺则 None（→ 不套，留 Word 默认）。"""
    return DU.pt_of(template, size) if (size is not None and size != "") else None


def _read_style(template, sec, *, page_break=None):
    """从模板某样式段读出排版规格——**只读模板，缺的留 None=不套**。
    键名做别名兼容（编译模板形状不一）；same_as 引用其他段。零自造格式值。
    """
    sec = sec if isinstance(sec, dict) else {}
    sa = sec.get("same_as")                       # same_as: heading_1 / body_paragraph
    if isinstance(sa, str):
        ref = (template.get("headings", {}) or {}).get(sa) if sa.startswith("level_") else None
        if ref is None and sa in ("body_paragraph", "body"):
            ref = template.get("body_paragraph")
        if isinstance(ref, dict):
            merged = dict(ref); merged.update({k: v for k, v in sec.items() if k != "same_as"})
            sec = merged
    font = sec.get("font")
    cn = latin = size = None
    if isinstance(font, dict):
        cn, latin, size = font.get("cn"), font.get("latin"), font.get("size")
    elif isinstance(font, str):
        cn = font
    if sec.get("font_cn") is not None:
        cn = sec["font_cn"]
    if sec.get("latin") is not None:
        latin = sec["latin"]
    if sec.get("size") is not None:
        size = sec["size"]
    ls = sec.get("line_spacing")
    return {
        "cn": cn, "latin": latin, "size": _opt_pt(template, size),
        "bold": sec.get("bold"),
        "align": sec.get("align", sec.get("alignment")),
        "first_line_chars": sec.get("first_line_indent_chars", sec.get("indent_chars")),
        "line_single": (str(ls) == "single") if ls is not None else None,
        "before_lines": sec.get("space_before_lines"),
        "after_lines": sec.get("space_after_lines"),
        "page_break": sec.get("page_break_before", sec.get("new_page", page_break)),
    }


# 结构标签 -> 模板里定义其格式的段落（格式值全来自 DeepSeek 抽的模板，引擎不自造）
def _spec_for(label, template):
    t = template
    CA = t.get("chinese_abstract", {}) or {}
    EA = t.get("english_abstract", {}) or {}
    TB = t.get("tables", {}) or {}
    FG = t.get("figures", {}) or {}
    REF = t.get("references", {}) or {}
    ACK = t.get("acknowledgements", {}) or {}
    TOC = t.get("table_of_contents", {}) or {}
    sec_map = {
        "body": t.get("body_paragraph"),
        "abstract_cn_title": CA.get("title"), "abstract_cn_body": CA.get("body"), "keywords_cn": CA.get("keywords"),
        "title_en": EA.get("title"), "abstract_en_title": EA.get("title"), "abstract_en_body": EA.get("body"),
        "keywords_en": EA.get("keywords"),
        "table_title": TB.get("title"), "figure_title": FG.get("title"), "note": TB.get("notes"),
        "reference_title": REF.get("title"), "reference_item": REF.get("body"),
        "ack_title": ACK.get("title"), "ack_body": ACK.get("body"),
        "toc_title": TOC.get("title"),
    }
    if label in ("heading_1", "heading_2", "heading_3", "heading_4", "heading_5"):
        return _read_style(t, (t.get("headings", {}) or {}).get(f"level_{label[-1]}"))
    if label in sec_map:
        return _read_style(t, sec_map[label])
    return _read_style(t, t.get("body_paragraph"))      # blank / 未知 -> 正文段样式


_LABEL_TO_SLOT = {"cover_doctype": "doctype", "title_main": "title_cn",
                  "title_en": "title_en", "cover_field": "field", "cover_date": "date"}


def _emit_cover(out, cover_blocks, template, labels):
    """封面：只读 template['cover']（DeepSeek 抽的）。logo + 槽位格式 + 空白行结构全来自模板，缺则不套。"""
    cv = template.get("cover", {}) or {}
    slots = cv.get("slots", {}) or {}
    blanks = cv.get("blanks", {}) or {}
    align = cv.get("align")
    bcfg = cv.get("blank", {}) if isinstance(cv.get("blank"), dict) else {}
    blank_cn = bcfg.get("font")
    blank_sz = _opt_pt(template, bcfg.get("size"))

    def slot_spec(slot):
        s = slots.get(slot) if isinstance(slots.get(slot), dict) else slots.get("other")
        spec = _read_style(template, s)
        if spec.get("align") is None:
            spec["align"] = align
        return spec

    def add_blanks(n):
        for _ in range(int(DU._num(n, 0))):
            p = out.add_paragraph()
            r = p.add_run(" ")
            DU.set_run_font(r, cn=blank_cn, size_pt=blank_sz)
            DU.set_paragraph_format(p, align=align, line_spacing_single=True)

    # 行1：校名标准字图（模板 furniture，无文字→不影响内容守恒）
    logo = cv.get("logo")
    if isinstance(logo, dict) and logo.get("asset"):
        asset = config.TEMPLATES_DIR / logo["asset"]
        if asset.exists():
            p = out.add_paragraph()
            p.alignment = DU.ALIGN["center"]
            try:
                kw = {}
                if logo.get("width_mm"):
                    kw["width"] = Mm(DU._num(logo["width_mm"], 120))
                if logo.get("height_mm"):
                    kw["height"] = Mm(DU._num(logo["height_mm"], 30))
                p.add_run().add_picture(str(asset), **kw)
            except Exception:
                pass

    inserted = set()
    for b in cover_blocks:
        if b["kind"] != "paragraph" or b.get("empty"):
            continue                                   # 丢源空行，按 blanks 自己控间距
        slot = _LABEL_TO_SLOT.get(labels.get(b["idx"], "cover"), "other")
        if slot in ("title_cn", "title_en") and "title" not in inserted:
            add_blanks(blanks.get("after_doctype", 0)); inserted.add("title")
        if slot == "field" and "field" not in inserted:
            add_blanks(blanks.get("after_title", 0)); inserted.add("field")
        if slot == "date" and "date" not in inserted:
            add_blanks(blanks.get("after_fields", 0)); inserted.add("date")
        _add_paragraph(out, b, slot_spec(slot))


def _add_paragraph(out, block, spec):
    p = out.add_paragraph()
    runs = block.get("runs") or ([{"text": block.get("text", ""), "bold": False, "italic": False}])
    for rd in runs:
        if rd["text"] == "":
            continue
        r = p.add_run(rd["text"])
        eff_bold = spec["bold"] if spec["bold"] is not None else rd["bold"]
        DU.set_run_font(r, cn=spec["cn"], latin=spec["latin"], size_pt=spec["size"],
                        bold=eff_bold, italic=rd["italic"])
    DU.set_paragraph_format(p, align=spec["align"], first_line_chars=spec["first_line_chars"],
                            line_spacing_single=spec["line_single"],
                            before_lines=spec.get("before_lines"), after_lines=spec.get("after_lines"),
                            page_break_before=spec.get("page_break"))
    return p


def _add_table(out, block, template):
    n_rows = max(block["n_rows"], 1); n_cols = max(block["n_cols"], 1)
    table = out.add_table(rows=n_rows, cols=n_cols)
    table.alignment = 1
    tconf = template.get("tables", {}) or {}
    cs = _read_style(template, tconf.get("content"))     # 内容字体只读模板
    for ri, row in enumerate(block["rows"]):
        for ci, cell_text in enumerate(row):
            if ri < n_rows and ci < n_cols:
                cell = table.cell(ri, ci)
                cell.text = cell_text
                for p in cell.paragraphs:
                    p.alignment = 1
                    for r in p.runs:
                        DU.set_run_font(r, cn=cs["cn"], latin=cs["latin"], size_pt=cs["size"])
    border = tconf.get("border", {}) if isinstance(tconf.get("border"), dict) else {}
    DU.apply_three_line_table(table, top_bottom_pt=DU._num(border.get("top_bottom_pt"), 1.5),
                              middle_pt=DU._num(border.get("middle_pt"), 0.5))
    return table


def _pick_title(blocks, labels):
    for b in blocks:
        if b["kind"] == "paragraph" and labels.get(b["idx"]) == "title_main" and b["text"].strip():
            return b["text"].strip()
    cands = [b["text"].strip() for b in blocks[:12] if b["kind"] == "paragraph" and b["text"].strip()]
    return max(cands, key=len) if cands else ""


def _partition(blocks, labels):
    """切成 封面 / 前置(摘要·目录) / 正文。"""
    n = len(blocks); cover_end = n; body_start = n
    for i, b in enumerate(blocks):
        lab = labels.get(b["idx"]) if b["kind"] == "paragraph" else None
        if cover_end == n and lab in _FRONT_MARKERS:
            cover_end = i
        if body_start == n and lab == "heading_1":
            body_start = i
    if body_start < cover_end:
        body_start = cover_end
    return blocks[:cover_end], blocks[cover_end:body_start], blocks[body_start:]


_TOC_PAGE_RE = re.compile(r"\s*([0-9IVXLCMivxlcm]+)\s*$")
_TOC_NUM_RE = re.compile(r"^\s*(\d+(?:\.\d+)*)")


def _add_toc_item(out, block, template):
    """目录条目（2.6）：章/节/条三级缩进 + 页码右对齐 + 点引线。数据来自 template['table_of_contents']。"""
    toc = template.get("table_of_contents", {}) or {}
    latin0 = toc.get("number_letter_font")
    text = block["text"].strip()
    page = ""
    title = text
    m = _TOC_PAGE_RE.search(text)
    if m:                                   # 末尾的数字/罗马 = 页码（剥掉只换 tab，_norm 比对仍守恒）
        page = m.group(1)
        title = text[:m.start()].rstrip()
    lm = _TOC_NUM_RE.match(title)
    level = min((lm.group(1).count(".") + 1) if lm else 1, 3)
    st = _read_style(template, toc.get(f"level_{level}"))     # 字体字号加粗只读模板
    cn = st["cn"]; latin = st["latin"] or latin0; bold = st["bold"]; size = st["size"]
    indent = st["first_line_chars"]
    if indent is None:
        indent = level - 1                  # 缩进缺则按层级递进（纯机制：目录靠层级缩进）

    p = out.add_paragraph()
    pg = template.get("page", {}) or {}
    mg = pg.get("margins_mm", {}) if isinstance(pg.get("margins_mm"), dict) else {}
    tw = DU._num(pg.get("width_mm"), 210) - DU._num(mg.get("left"), 25.4) - DU._num(mg.get("right"), 25.4)
    p.paragraph_format.tab_stops.add_tab_stop(Mm(tw), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r1 = p.add_run(title + ("\t" if page else ""))
    DU.set_run_font(r1, cn=cn, latin=latin, size_pt=size, bold=bold)
    if page:
        r2 = p.add_run(page)
        DU.set_run_font(r2, cn=cn, latin=latin, size_pt=size, bold=False)   # 页码不加粗
    DU.set_paragraph_format(p, align="left", first_line_chars=indent, line_spacing_single=True)


def _emit(out, blocks, template, labels=None):
    """前置/正文区：按结构标签查 _spec_for 套版。"""
    for b in blocks:
        if b["kind"] == "table":
            _add_table(out, b, template)
            continue
        lab = labels.get(b["idx"], "body") if labels else "body"
        if lab == "toc_item":
            _add_toc_item(out, b, template)
            continue
        _add_paragraph(out, b, _spec_for(lab, template))


def format_docx(blocks, labels, template, out_path):
    out = Document()
    # 全局清场：去掉 Word 默认的段后间距（Office 风格默认，非格式判断），让模板/源说了算
    nf = out.styles["Normal"].paragraph_format
    nf.space_before = Pt(0); nf.space_after = Pt(0)
    title = _pick_title(blocks, labels)
    cover, front, body = _partition(blocks, labels)

    groups = [(cover, "cover"), (front, "front"), (body, "body")]
    groups = [(b, role) for (b, role) in groups if b]
    if not groups:
        groups = [(blocks, "body")]

    roles = []
    for i, (blks, role) in enumerate(groups):
        if i > 0:
            out.add_section(WD_SECTION.NEW_PAGE)
        if role == "cover":
            _emit_cover(out, blks, template, labels)
        else:
            _emit(out, blks, template, labels)
        roles.append(role)

    # 删掉 Document() 自带的首个空段（封面顶部多余空行）
    ps = out.paragraphs
    if ps and not ps[0].text.strip() and not ps[0].runs:
        ps[0]._element.getparent().remove(ps[0]._element)

    # 逐节配置页面 + 页眉 + 分节页码（页码格式读 template['pagination']）
    pag = template.get("pagination", {}) or {}
    NUMFMT = {"roman": "lowerRoman", "lower_roman": "lowerRoman", "lowerroman": "lowerRoman",
              "upper_roman": "upperRoman", "arabic": "decimal", "decimal": "decimal"}
    fp = (template.get("footer", {}) or {}).get("page_number", {})
    foot_fmt = (fp.get("format") if isinstance(fp, dict) else None) or "n"
    for sec, role in zip(out.sections, roles):
        DU.set_page(sec, template)
        if role == "cover":
            DU.blank_header_footer(sec)                              # 封面/扉页：无页眉、无页码
        else:
            DU.setup_title_header(sec, template, title)
            key = "frontmatter" if role == "front" else "body"
            numfmt = NUMFMT.get(str(pag.get(key)).lower(), "lowerRoman" if role == "front" else "decimal")
            DU.setup_pagenum_footer(sec, template, "n" if role == "front" else foot_fmt)
            DU.set_page_number_format(sec, numfmt, start=1)

    out.save(out_path)
    return {"out_path": out_path, "title": title, "change_log": _change_log(template),
            "sections": roles}


def _change_log(t):
    """大白话改动清单（按模板实际值，不写死任何格式判断）。"""
    meta = t.get("meta", {}) or {}
    log = [{"what": "套用规范", "detail": "《%s》" % (meta.get("name") or "目标规范")}]
    pg = t.get("page", {}) or {}
    m = pg.get("margins_mm") if isinstance(pg.get("margins_mm"), dict) else None
    if m:
        log.append({"what": "页边距", "detail": "上%s/下%s/左%s/右%s mm" % (m.get("top"), m.get("bottom"), m.get("left"), m.get("right"))})
    log.append({"what": "结构套版", "detail": "封面/标题分级/正文/摘要/目录/参考文献等按规范逐元素套用（内容一字不动）"})
    log.append({"what": "分节页码", "detail": "封面/扉页不编 → 前置/正文按规范页码格式"})
    log.append({"what": "三线表 + 页眉", "detail": "表格三线格 + 正文区页眉"})
    return log
