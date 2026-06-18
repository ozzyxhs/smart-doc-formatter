"""建设期 LLM 缝②：规范文档 → 声明式模板 YAML（"规范→模板"编译器）。

两种输入都支持：
  - 规范条文文档(pdf/docx 里是规则文字) → LLM 读规则；
  - 样本/模板 docx → python-docx 读「实际生效的格式」(边距/样式/字体/字号) → LLM 归并成 YAML。
样本模式更可靠：格式数值来自文件本身，不靠 LLM 猜。
"""
import re
import json
import uuid
import subprocess
from pathlib import Path
import yaml
from docx import Document
from docx.oxml.ns import qn
from . import llm, config


def _emu_mm(v):
    return round(v / 36000, 1) if v else None


def _run_font(run):
    rPr = run._element.find(qn("w:rPr"))
    ea = lat = sz = None
    b = run.bold
    if rPr is not None:
        rf = rPr.find(qn("w:rFonts"))
        if rf is not None:
            ea = rf.get(qn("w:eastAsia")); lat = rf.get(qn("w:ascii"))
        s = rPr.find(qn("w:sz"))
        if s is not None:
            sz = int(s.get(qn("w:val"))) / 2
    return ea, lat, sz, b


def _docx_facts(doc):
    """读样本 docx 里实际生效的格式：页面/边距 + 各样式的字体字号缩进样本。"""
    sec = doc.sections[0]
    facts = {
        "page": {"width_mm": _emu_mm(sec.page_width), "height_mm": _emu_mm(sec.page_height),
                 "margins_mm": {"top": _emu_mm(sec.top_margin), "bottom": _emu_mm(sec.bottom_margin),
                                "left": _emu_mm(sec.left_margin), "right": _emu_mm(sec.right_margin)}},
        "styles": {},
    }
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        sname = (p.style.name if p.style else "Normal") or "Normal"
        st = facts["styles"].setdefault(sname, {"count": 0, "samples": []})
        st["count"] += 1
        if len(st["samples"]) < 3 and p.runs:
            ea, lat, sz, b = _run_font(p.runs[0])
            fli = p.paragraph_format.first_line_indent
            st["samples"].append({"text": t[:28], "eastAsia": ea, "latin": lat, "pt": sz, "bold": b,
                                  "align": str(p.alignment), "first_line_indent_mm": _emu_mm(fli) if fli else None})
    return facts


def _doc_to_docx(src, dst):
    script = ("$ErrorActionPreference='Stop'\n$w=New-Object -ComObject Word.Application\n$w.Visible=$false\n"
              f"try{{$d=$w.Documents.Open([string]'{src}',$false,$true);$d.SaveAs([string]'{dst}',[int]16);$d.Close($false)}}finally{{$w.Quit()}}")
    ps1 = Path(dst).with_suffix(".conv.ps1")
    ps1.write_text(script, encoding="utf-8-sig")
    try:
        subprocess.run(["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-File", str(ps1)],
                       capture_output=True, timeout=120)
    finally:
        try:
            ps1.unlink()
        except OSError:
            pass


def _read_spec(path):
    path = Path(path)
    ext = path.suffix.lower()
    text, facts = "", None
    if ext == ".doc":
        dst = path.with_suffix(".converted.docx")
        _doc_to_docx(path, dst)
        path, ext = dst, ".docx"
    if ext == ".docx":
        d = Document(str(path))
        text = "\n".join(p.text for p in d.paragraphs if p.text.strip())[:6000]
        facts = _docx_facts(d)
    elif ext == ".pdf":
        from pypdf import PdfReader
        r = PdfReader(str(path))
        text = "\n".join((pg.extract_text() or "") for pg in r.pages)[:8000]
    return text, facts


def _parse_yaml(text):
    t = (text or "").strip()
    t = re.sub(r"^```(?:yaml)?\s*|\s*```$", "", t, flags=re.S)
    return yaml.safe_load(t)


_SYS = (
    "你是排版规范编译器。把给定的【规范/样本】抽成一份与【schema 样例】完全同结构的 YAML 模板，"
    "供确定性排版引擎读取。规则：①优先采用【样本实际格式 facts】里的真实数值(边距/字号pt/字体/缩进)；"
    "②规范文本里有明确规则的也采纳；③字号尽量转成中文字号名(小二/三号/五号…)，并保留 size_table；"
    "④封面/摘要/目录/标题分级/正文/三线表/页眉页脚/页码/参考文献等字段尽量补全，缺的给合理默认；"
    "⑤meta.id/name/institution/doc_type 留占位，调用方会覆盖。只输出 YAML 本体，不要解释、不要代码围栏。"
)


def _extract_yaml(text, facts, *, name, institution, doc_type):
    from . import template_norm
    example = yaml.safe_dump(template_norm.DEFAULTS, allow_unicode=True, sort_keys=False)
    user = (
        f"【schema 样例（引擎标准模板结构，仅示意结构，不要照抄数值）】\n{example}\n\n"
        f"【新规范】name={name} institution={institution} doc_type={doc_type}\n\n"
        f"【样本实际格式 facts（来自文件，最可信）】\n{json.dumps(facts, ensure_ascii=False) if facts else '（无，纯文本规范）'}\n\n"
        f"【规范/样本文本】\n{text}"
    )
    out = llm.chat([{"role": "system", "content": _SYS}, {"role": "user", "content": user}],
                   thinking=True, max_tokens=12000)
    y = _parse_yaml(out) or {}
    y.setdefault("meta", {})
    y["meta"].update({"name": name, "institution": institution, "doc_type": doc_type})
    return y


def extract_template(path, *, tid, name, institution, doc_type="thesis"):
    """规范文档 → draft YAML(dict) + facts（供 CLI/测试）。"""
    text, facts = _read_spec(path)
    y = _extract_yaml(text, facts, name=name, institution=institution, doc_type=doc_type)
    y["meta"]["id"] = tid
    return y, facts


_CHECK_SYS = (
    "你是规范抽取的算力自检员（0 人工）。给你【抽出的模板 YAML】和【规范/样本来源】。"
    "把 YAML 的关键格式规则逐条'回译'成人话，与来源交叉核对，判定每条状态："
    "ok=来源明确支持；infer=合理推断/需报备；missing=来源没说、用了默认、待补；wrong=与来源矛盾。"
    "只输出 JSON：{\"ok\": true/false, \"items\": [{\"field\":\"\",\"yaml\":\"\",\"source\":\"\",\"status\":\"ok|infer|missing|wrong\",\"note\":\"\"}]}。"
    "出现任一 wrong 则 ok=false。"
)


def self_check(yaml_obj, text, facts=None):
    """算力自检·回译交叉校验：YAML 是否被来源支持。返回 {ok, items, error?}。"""
    user = ("【抽出的模板 YAML】\n" + yaml.safe_dump(yaml_obj, allow_unicode=True, sort_keys=False)[:8000]
            + "\n\n【来源 facts】\n" + (json.dumps(facts, ensure_ascii=False) if facts else "（无）")
            + "\n\n【来源文本】\n" + (text or "")[:6000])
    try:
        out = llm.chat_json([{"role": "system", "content": _CHECK_SYS}, {"role": "user", "content": user}],
                            thinking=True, max_tokens=16000)   # 思考模式 COT+输出共用，给足避免 JSON 截断
        out.setdefault("items", [])
        out.setdefault("ok", not any(i.get("status") == "wrong" for i in out["items"]))
    except Exception as e:
        return {"ok": None, "items": [], "error": str(e)}
    return out


def _slug(doc_type):
    return f"custom-{doc_type}-{uuid.uuid4().hex[:6]}"


def _summary(y):
    """人话版抽取摘要（给"上传新规范"屏的卡片）。"""
    out = []
    pg = y.get("page", {}).get("margins_mm")
    if pg:
        out.append({"k": "页边距", "v": f"上{pg.get('top')}/下{pg.get('bottom')}/左{pg.get('left')}/右{pg.get('right')} mm"})
    bp = y.get("body_paragraph", {}).get("font", {})
    if bp:
        out.append({"k": "正文字体", "v": f"{bp.get('cn')} {bp.get('size')} / 西文 {bp.get('latin')}"})
    h1 = y.get("headings", {}).get("level_1", {})
    if h1:
        out.append({"k": "一级标题", "v": f"{h1.get('font_cn')} {h1.get('size')}"})
    nlv = len([k for k in y.get("headings", {}) if k.startswith("level_")])
    if nlv:
        out.append({"k": "标题层级", "v": f"{nlv} 级"})
    refs = y.get("references", {}).get("requirements", {})
    if refs:
        out.append({"k": "参考文献", "v": json.dumps(refs, ensure_ascii=False)})
    cov = y.get("cover", {}).get("slots")
    if cov:
        out.append({"k": "封面要素", "v": "、".join(list(cov.keys()))})
    return out


def compile_and_save(path, *, name, institution, doc_type="thesis"):
    """上传规范 → 抽取 + 自检 + 入库。返回 {tid, name, summary, audit}。"""
    text, facts = _read_spec(path)
    y = _extract_yaml(text, facts, name=name, institution=institution, doc_type=doc_type)
    tid = _slug(doc_type)
    y["meta"]["id"] = tid
    audit = self_check(y, text, facts)
    (config.TEMPLATES_DIR / f"{tid}.yaml").write_text(
        yaml.safe_dump(y, allow_unicode=True, sort_keys=False), encoding="utf-8")
    return {"tid": tid, "name": name, "summary": _summary(y), "audit": audit}
