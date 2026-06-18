"""C0 测试用的合成最小学位论文 docx —— 代码生成、不入库二进制，确保可复现。

覆盖今天确实存在的路径：封面槽位 / 摘要(前置区边界) / 一级标题 / 正文 / 三线表 / 参考文献。
TOC、封面精确排布等留待各自元素 issue 落地时附自己的 golden（金标随元素长）。
"""
from docx import Document

# (段落文本, 离线分类标签) —— 按文档顺序；标签驱动 offline_labels 打桩分类器。
BLOCKS = [
    ("黑龙江某某大学", "cover_field"),
    ("本科毕业论文", "cover_doctype"),
    ("一个用于离线测试的最小合成标题", "title_main"),
    ("摘要", "abstract_cn_title"),
    ("这是摘要正文，仅用于验证排版前后内容守恒。", "abstract_cn_body"),
    ("关键词：测试；排版；内容守恒", "keywords_cn"),
    ("1 引言", "heading_1"),
    ("这是正文第一段，包含若干中文字符用于字数统计与守恒比对。", "body"),
    ("这是正文第二段，继续验证逐段拷贝不丢字。", "body"),
    ("参考文献", "reference_title"),
    ("[1] 张三. 一种测试方法[J]. 测试学报, 2026, 1(1): 1-2.", "reference_item"),
]

_TABLE_AFTER = BLOCKS[7][0]          # 表插在正文第一段之后（确保落在正文区、不在封面区）
_TABLE_ROWS = [["列A", "列B"], ["甲", "乙"], ["丙", "丁"]]


def build_synthetic_thesis(path):
    """生成最小合成论文 docx 到 path。返回 {段落文本: 标签} 供分类打桩。"""
    doc = Document()
    for text, _label in BLOCKS:
        doc.add_paragraph(text)
        if text == _TABLE_AFTER:
            t = doc.add_table(rows=len(_TABLE_ROWS), cols=len(_TABLE_ROWS[0]))
            for r, row in enumerate(_TABLE_ROWS):
                for c, val in enumerate(row):
                    t.cell(r, c).text = val
    doc.save(str(path))
    return {text: label for text, label in BLOCKS}


def offline_labels(blocks):
    """按段落文本映射回标签（离线、确定性，替代真 LLM 分类）。"""
    t2l = {text: label for text, label in BLOCKS}
    return {b["idx"]: t2l.get(b["text"].strip(), "body")
            for b in blocks if b["kind"] == "paragraph"}


def build_min_with_heading_style(path):
    """带真 Word「Heading 1」样式的最小 docx —— 供 C1 测试『LLM 失败但样式兜底可用』路径。"""
    doc = Document()
    doc.add_paragraph("这是正文第一段，用于守恒比对。")
    doc.add_paragraph("1 引言", style="Heading 1")
    doc.add_paragraph("这是正文第二段，继续比对。")
    doc.save(str(path))


def build_min_unstyled(path):
    """全 Normal 样式、无标题样式可兜底 —— 供 C1 测试『兜底无用必须阻断』路径。"""
    doc = Document()
    doc.add_paragraph("第一段纯正文，无任何标题样式。")
    doc.add_paragraph("第二段纯正文，继续比对内容守恒。")
    doc.save(str(path))
